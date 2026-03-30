from copy import deepcopy
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO
from typing import Dict, List

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
class DocumentGenerator:
    @staticmethod
    def _apply_times_new_roman(run) -> None:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:ascii"), "Times New Roman")
        r_fonts.set(qn("w:hAnsi"), "Times New Roman")
        r_fonts.set(qn("w:eastAsia"), "Times New Roman")
        r_fonts.set(qn("w:cs"), "Times New Roman")

    @staticmethod
    def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)
        if not full_text:
            return

        replacements: list[tuple[int, int, str]] = []
        for key, value in mapping.items():
            start = 0
            while True:
                index = full_text.find(key, start)
                if index == -1:
                    break
                replacements.append((index, index + len(key), value))
                start = index + len(key)

        if not replacements:
            return

        replacements.sort(key=lambda item: item[0], reverse=True)

        run_ranges = []
        cursor = 0
        for run in paragraph.runs:
            run_text = run.text
            run_start = cursor
            run_end = cursor + len(run_text)
            run_ranges.append((run, run_start, run_end))
            cursor = run_end

        for start, end, replacement in replacements:
            affected = [
                (run, run_start, run_end)
                for run, run_start, run_end in run_ranges
                if run_end > start and run_start < end
            ]
            if not affected:
                continue

            first_run, first_start, first_end = affected[0]
            last_run, last_start, last_end = affected[-1]

            prefix = first_run.text[: max(0, start - first_start)]
            suffix = last_run.text[max(0, end - last_start):]
            first_run.text = f"{prefix}{replacement}{suffix}"
            DocumentGenerator._apply_times_new_roman(first_run)

            for run, _, _ in affected[1:]:
                run.text = ""

            delta = len(replacement) - (end - start)
            updated_ranges = []
            for run, run_start, run_end in run_ranges:
                if run_end <= start:
                    updated_ranges.append((run, run_start, run_end))
                elif run_start >= end:
                    updated_ranges.append((run, run_start + delta, run_end + delta))
                elif run is first_run:
                    new_end = first_start + len(first_run.text)
                    updated_ranges.append((run, first_start, new_end))
                else:
                    updated_ranges.append((run, run_start, run_start))
            run_ranges = updated_ranges

    @staticmethod
    def _replace_placeholders(doc: Document, mapping: Dict[str, str]) -> None:
        for paragraph in doc.paragraphs:
            DocumentGenerator._replace_in_paragraph(paragraph, mapping)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        DocumentGenerator._replace_in_paragraph(paragraph, mapping)

    @staticmethod
    def _replace_row_placeholders(row, mapping: Dict[str, str]) -> None:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                DocumentGenerator._replace_in_paragraph(paragraph, mapping)

    @staticmethod
    def _save(doc: Document) -> BytesIO:
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_from_template(template_path: str, mapping: Dict[str, str]) -> BytesIO:
        doc = Document(template_path)
        DocumentGenerator._replace_placeholders(doc, mapping)
        return DocumentGenerator._save(doc)

    @staticmethod
    def to_decimal(value: float | str | Decimal) -> Decimal:
        return Decimal(str(value)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)

    @staticmethod
    def format_number(value: float | str | Decimal) -> str:
        normalized = DocumentGenerator.to_decimal(value)
        as_string = format(normalized, "f")
        if "." in as_string:
            as_string = as_string.rstrip("0").rstrip(".")
        return as_string.replace(".", ",")

    @staticmethod
    def _format_decimal(value: Decimal) -> str:
        normalized = value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        as_string = format(normalized, "f")
        if "." in as_string:
            as_string = as_string.rstrip("0").rstrip(".")
        return as_string.replace(".", ",")

    @staticmethod
    def _plural(number: int, forms: tuple[str, str, str]) -> str:
        number = abs(number) % 100
        last = number % 10
        if 10 < number < 20:
            return forms[2]
        if 1 < last < 5:
            return forms[1]
        if last == 1:
            return forms[0]
        return forms[2]

    @staticmethod
    def _number_to_words(number: int, female: bool = False) -> str:
        units = {
            False: ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"],
            True: ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"],
        }
        teens = [
            "десять",
            "одиннадцать",
            "двенадцать",
            "тринадцать",
            "четырнадцать",
            "пятнадцать",
            "шестнадцать",
            "семнадцать",
            "восемнадцать",
            "девятнадцать",
        ]
        tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
        hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот", "шестьсот", "семьсот", "восемьсот", "девятьсот"]
        orders = [
            (("", "", ""), False),
            (("тысяча", "тысячи", "тысяч"), True),
            (("миллион", "миллиона", "миллионов"), False),
            (("миллиард", "миллиарда", "миллиардов"), False),
        ]

        if number == 0:
            return "ноль"

        parts: List[str] = []
        order_index = 0
        remaining = number

        while remaining > 0:
            chunk = remaining % 1000
            remaining //= 1000

            if chunk:
                chunk_parts: List[str] = []
                chunk_hundreds = chunk // 100
                chunk_tens = (chunk % 100) // 10
                chunk_units = chunk % 10

                if chunk_hundreds:
                    chunk_parts.append(hundreds[chunk_hundreds])

                if chunk_tens == 1:
                    chunk_parts.append(teens[chunk_units])
                else:
                    if chunk_tens:
                        chunk_parts.append(tens[chunk_tens])
                    if chunk_units:
                        use_female = orders[order_index][1]
                        chunk_parts.append(units[use_female][chunk_units])

                order_forms = orders[order_index][0]
                if order_forms[0]:
                    chunk_parts.append(DocumentGenerator._plural(chunk, order_forms))

                parts.insert(0, " ".join(part for part in chunk_parts if part))

            order_index += 1

        return " ".join(part for part in parts if part)

    @staticmethod
    def format_rub_amount_text(value: float | Decimal) -> str:
        amount = DocumentGenerator.to_decimal(value)
        rubles = int(amount)
        kopeks = int((amount - Decimal(rubles)) * 100)
        rubles_text = DocumentGenerator._number_to_words(rubles)
        rubles_unit = DocumentGenerator._plural(rubles, ("рубль", "рубля", "рублей"))
        kopeks_unit = DocumentGenerator._plural(kopeks, ("копейка", "копейки", "копеек"))
        return f"{rubles_text} {rubles_unit} {kopeks:02d} {kopeks_unit}"

    @staticmethod
    def generate_payment_memo_from_template(
        template_path: str,
        mapping: Dict[str, str],
        payment_rows: List[Dict[str, str]],
        total_hours: str,
        total_amount: str,
        total_amount_text: str,
    ) -> BytesIO:
        doc = Document(template_path)
        DocumentGenerator._replace_placeholders(doc, mapping)

        if len(doc.tables) < 2:
            raise ValueError("Payment template does not contain the expected payments table")

        payments_table = doc.tables[1]
        if len(payments_table.rows) < 3:
            raise ValueError("Payment template table must contain a header, template row, and total row")

        template_row = payments_table.rows[1]
        total_row = payments_table.rows[2]
        template_row_xml = deepcopy(template_row._tr)

        DocumentGenerator._replace_row_placeholders(template_row, payment_rows[0])

        for row_mapping in payment_rows[1:]:
            new_row = deepcopy(template_row_xml)
            total_row._tr.addprevious(new_row)
            inserted_row = payments_table.rows[-2]
            DocumentGenerator._replace_row_placeholders(inserted_row, row_mapping)

        DocumentGenerator._replace_row_placeholders(
            total_row,
            {
                "{{HOURS}}": total_hours,
                "{{SUM}}": total_amount,
            },
        )

        DocumentGenerator._replace_placeholders(
            doc,
            {
                "{{SUM}}": total_amount,
                "{{SUM_TEXT}}": total_amount_text,
            },
        )

        return DocumentGenerator._save(doc)
